import os
import subprocess
from .base_converter import BaseConverter

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

class CodeConverter(BaseConverter):
    def __init__(self):
        super().__init__()
        self.supported_input_formats = [
            'py', 'java', 'cpp', 'c', 'js', 'html', 'css', 'php', 'rb', 'go',
            'ts', 'jsx', 'tsx', 'cs', 'vb', 'sql', 'xml', 'json', 'yaml', 'yml'
        ]
        self.supported_output_formats = [
            'pdf', 'docx', 'html', 'txt'
        ]

        # Syntax highlighting mappings for HTML output
        self.language_mappings = {
            'py': 'python',
            'js': 'javascript',
            'ts': 'typescript',
            'jsx': 'javascript',
            'tsx': 'typescript',
            'cpp': 'cpp',
            'c': 'c',
            'cs': 'csharp',
            'java': 'java',
            'php': 'php',
            'rb': 'ruby',
            'go': 'go',
            'html': 'html',
            'css': 'css',
            'xml': 'xml',
            'json': 'json',
            'yaml': 'yaml',
            'yml': 'yaml',
            'sql': 'sql'
        }

    def convert(self, input_path, output_format, output_directory):
        """Convert code files to documentation formats."""
        try:
            if not self.ensure_output_directory(output_directory):
                return False

            output_path = self.get_output_filename(input_path, output_format, output_directory)

            if output_format.lower() == 'txt':
                return self._convert_to_txt(input_path, output_path)
            elif output_format.lower() == 'html':
                return self._convert_to_html(input_path, output_path)
            elif output_format.lower() == 'docx':
                return self._convert_to_docx(input_path, output_path)
            elif output_format.lower() == 'pdf':
                return self._convert_to_pdf(input_path, output_path)

            return False

        except Exception as e:
            print(f"Code conversion error: {str(e)}")
            return False

    def _convert_to_txt(self, input_path, output_path):
        """Convert code to formatted text file."""
        try:
            with open(input_path, 'r', encoding='utf-8') as infile:
                content = infile.read()

            # Add header with file information
            header = f"Source Code Documentation\n"
            header += f"File: {os.path.basename(input_path)}\n"
            header += f"Language: {self.get_file_extension(input_path).upper()}\n"
            header += f"{'=' * 50}\n\n"

            with open(output_path, 'w', encoding='utf-8') as outfile:
                outfile.write(header)
                outfile.write(content)

            return True

        except Exception as e:
            print(f"TXT conversion error: {str(e)}")
            return False

    def _convert_to_html(self, input_path, output_path):
        """Convert code to HTML with syntax highlighting."""
        try:
            with open(input_path, 'r', encoding='utf-8') as infile:
                content = infile.read()

            file_ext = self.get_file_extension(input_path)
            language = self.language_mappings.get(file_ext, 'text')

            html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Source Code: {os.path.basename(input_path)}</title>
    <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/highlight.js/11.8.0/styles/default.min.css">
    <script src="https://cdnjs.cloudflare.com/ajax/libs/highlight.js/11.8.0/highlight.min.js"></script>
    <style>
        body {{
            font-family: Arial, sans-serif;
            margin: 40px;
            background-color: #f5f5f5;
        }}
        .header {{
            background-color: #333;
            color: white;
            padding: 20px;
            border-radius: 8px;
            margin-bottom: 20px;
        }}
        .code-container {{
            background-color: white;
            border-radius: 8px;
            overflow: hidden;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }}
        pre {{
            margin: 0;
            padding: 20px;
            overflow-x: auto;
        }}
        code {{
            font-family: 'Courier New', monospace;
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>Source Code Documentation</h1>
        <p><strong>File:</strong> {os.path.basename(input_path)}</p>
        <p><strong>Language:</strong> {language.title()}</p>
    </div>
    <div class="code-container">
        <pre><code class="language-{language}">{self._escape_html(content)}</code></pre>
    </div>
    <script>hljs.highlightAll();</script>
</body>
</html>"""

            with open(output_path, 'w', encoding='utf-8') as outfile:
                outfile.write(html_content)

            return True

        except Exception as e:
            print(f"HTML conversion error: {str(e)}")
            return False

    def _convert_to_docx(self, input_path, output_path):
        """Convert code to DOCX document."""
        if not PYTHON_DOCX_AVAILABLE:
            print("python-docx not available for DOCX conversion")
            return False

        try:
            with open(input_path, 'r', encoding='utf-8') as infile:
                content = infile.read()

            doc = Document()

            # Add title
            title = doc.add_heading('Source Code Documentation', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Add file information
            info_para = doc.add_paragraph()
            info_para.add_run('File: ').bold = True
            info_para.add_run(os.path.basename(input_path))
            info_para.add_run('\nLanguage: ').bold = True
            info_para.add_run(self.get_file_extension(input_path).upper())

            # Add separator
            doc.add_paragraph('_' * 50)

            # Add code content
            code_para = doc.add_paragraph()
            code_run = code_para.add_run(content)
            code_run.font.name = 'Courier New'
            code_run.font.size = Inches(0.1)

            doc.save(output_path)
            return True

        except Exception as e:
            print(f"DOCX conversion error: {str(e)}")
            return False

    def _convert_to_pdf(self, input_path, output_path):
        """Convert code to PDF via HTML intermediate."""
        try:
            # First convert to HTML
            temp_html = output_path.replace('.pdf', '_temp.html')
            if not self._convert_to_html(input_path, temp_html):
                return False

            # Try to convert HTML to PDF using wkhtmltopdf
            if self._convert_html_to_pdf_wkhtmltopdf(temp_html, output_path):
                os.remove(temp_html)
                return True

            # Fallback: try converting via pandoc
            if self._convert_html_to_pdf_pandoc(temp_html, output_path):
                os.remove(temp_html)
                return True

            # Clean up temp file
            if os.path.exists(temp_html):
                os.remove(temp_html)

            return False

        except Exception as e:
            print(f"PDF conversion error: {str(e)}")
            return False

    def _convert_html_to_pdf_wkhtmltopdf(self, html_path, pdf_path):
        """Convert HTML to PDF using wkhtmltopdf."""
        try:
            cmd = ['wkhtmltopdf', '--page-size', 'A4', html_path, pdf_path]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
            return result.returncode == 0 and os.path.exists(pdf_path)
        except:
            return False

    def _convert_html_to_pdf_pandoc(self, html_path, pdf_path):
        """Convert HTML to PDF using pandoc."""
        try:
            cmd = ['pandoc', html_path, '-o', pdf_path]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
            return result.returncode == 0 and os.path.exists(pdf_path)
        except:
            return False

    def _escape_html(self, text):
        """Escape HTML special characters."""
        return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

    def analyze_code(self, file_path):
        """Analyze code file and return statistics."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                lines = f.readlines()

            total_lines = len(lines)
            code_lines = 0
            comment_lines = 0
            blank_lines = 0

            # Simple analysis based on file extension
            file_ext = self.get_file_extension(file_path)
            comment_patterns = {
                'py': ['#'],
                'java': ['//', '/*', '*'],
                'cpp': ['//', '/*', '*'],
                'c': ['//', '/*', '*'],
                'js': ['//', '/*', '*'],
                'css': ['/*', '*'],
                'html': ['<!--'],
                'php': ['//', '#', '/*', '*']
            }

            patterns = comment_patterns.get(file_ext, ['#', '//'])

            for line in lines:
                stripped = line.strip()
                if not stripped:
                    blank_lines += 1
                elif any(stripped.startswith(p) for p in patterns):
                    comment_lines += 1
                else:
                    code_lines += 1

            return {
                'total_lines': total_lines,
                'code_lines': code_lines,
                'comment_lines': comment_lines,
                'blank_lines': blank_lines,
                'file_size': os.path.getsize(file_path)
            }

        except Exception as e:
            print(f"Code analysis error: {str(e)}")
            return None